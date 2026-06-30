import unittest
from unittest.mock import patch, MagicMock
import os
import sys
import pandas as pd
import docx
import dotenv

# パス追加
script_dir = os.path.dirname(os.path.abspath(__file__))
dotenv.load_dotenv(os.path.abspath(os.path.join(script_dir, '../.env')))
sys.path.append(os.path.abspath(os.path.join(script_dir, '../scripts')))

import document_parser

class TestDocumentParser(unittest.TestCase):
    def setUp(self):
        self.test_dir = os.path.join(script_dir, "parser_test_tmp")
        os.makedirs(self.test_dir, exist_ok=True)
        
        self.docx_path = os.path.join(self.test_dir, "test_doc.docx")
        self.excel_path = os.path.join(self.test_dir, "test_sheet.xlsx")
        
        # 1. ダミーの Word ドキュメント作成
        doc = docx.Document()
        doc.add_heading("Main Document Title", level=1)
        doc.add_paragraph("This is a simple paragraph.")
        
        # テーブル作成
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header A"
        table.cell(0, 1).text = "Header B"
        table.cell(1, 0).text = "Data 1"
        table.cell(1, 1).text = "Data 2"
        
        doc.save(self.docx_path)
        
        # 2. ダミーの Excel ドキュメント作成
        df = pd.DataFrame({
            "Name": ["Alice", "Bob"],
            "Age": [25, 30]
        })
        with pd.ExcelWriter(self.excel_path) as writer:
            df.to_excel(writer, sheet_name="Employees", index=False)

    def tearDown(self):
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
        if os.path.exists(self.excel_path):
            os.remove(self.excel_path)
        if os.path.exists(self.test_dir):
            os.rmdir(self.test_dir)

    def test_parse_docx_to_markdown(self):
        """Test Word parser preserves structure, headings, and tables"""
        result = document_parser.parse_docx_to_markdown(self.docx_path)
        
        self.assertIn("# Main Document Title", result)
        self.assertIn("This is a simple paragraph.", result)
        self.assertIn("| Header A | Header B |", result)
        self.assertIn("| Data 1 | Data 2 |", result)

    def test_parse_excel_to_markdown(self):
        """Test Excel parser converts sheet to pandas markdown representation"""
        result = document_parser.parse_excel_to_markdown(self.excel_path)
        
        self.assertIn("## Sheet: Employees", result)
        self.assertIn("| Name   |   Age |", result)
        self.assertIn("| Alice  |    25 |", result)
        self.assertIn("| Bob    |    30 |", result)

    @patch('pypdf.PdfReader')
    def test_parse_pdf_to_markdown_native_text(self, mock_pdf_reader):
        """Test PDF parser with native text skips Vision and extracts text directly"""
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page content with enough characters to avoid vision parser."
        
        mock_reader_inst = MagicMock()
        mock_reader_inst.pages = [mock_page]
        mock_pdf_reader.return_value = mock_reader_inst
        
        # pdf2image は呼ばれないことを保証するため patch でエラーにするか監視
        with patch('pdf2image.convert_from_path') as mock_convert:
            result = document_parser.parse_pdf_to_markdown("dummy.pdf", use_vision=True)
            self.assertIn("Page content with enough characters", result)
            mock_convert.assert_not_called()

    @patch('pypdf.PdfReader')
    @patch('pdf2image.convert_from_path')
    @patch('document_parser.analyze_image_with_vision')
    def test_parse_pdf_to_markdown_vision_scanned(self, mock_vision, mock_convert, mock_pdf_reader):
        """Test PDF parser with empty/scanned page triggers Vision analysis"""
        mock_page = MagicMock()
        # テキストが少ない（または空）
        mock_page.extract_text.return_value = "   "
        
        mock_reader_inst = MagicMock()
        mock_reader_inst.pages = [mock_page]
        mock_pdf_reader.return_value = mock_reader_inst
        
        mock_image = MagicMock()
        mock_convert.return_value = [mock_image]
        mock_vision.return_value = "Optimized Vision Markdown Text"
        
        result = document_parser.parse_pdf_to_markdown("scanned.pdf", use_vision=True)
        
        # 変換・Visionコールが実行されたことを検証
        mock_convert.assert_called_once_with("scanned.pdf", first_page=1, last_page=1)
        mock_vision.assert_called_once_with(mock_image)
        
        self.assertIn("Vision OCR Result", result)
        self.assertIn("Optimized Vision Markdown Text", result)

    def test_routing_routing_by_extension(self):
        """Test convert_document_to_markdown routes correctly based on file suffix"""
        with patch('document_parser.parse_docx_to_markdown') as mock_docx, \
             patch('document_parser.parse_pdf_to_markdown') as mock_pdf, \
             patch('document_parser.parse_excel_to_markdown') as mock_xlsx, \
             patch('document_parser.parse_image_to_markdown') as mock_image:
             
             document_parser.convert_document_to_markdown("hello.docx")
             mock_docx.assert_called_once_with("hello.docx")
             
             document_parser.convert_document_to_markdown("hello.pdf")
             mock_pdf.assert_called_once_with("hello.pdf")
             
             document_parser.convert_document_to_markdown("hello.xlsx")
             mock_xlsx.assert_called_once_with("hello.xlsx")
             
             document_parser.convert_document_to_markdown("hello.png")
             mock_image.assert_called_once_with("hello.png")

    @patch('PIL.Image.open')
    @patch('document_parser.analyze_image_with_vision')
    def test_parse_image_to_markdown(self, mock_vision, mock_open):
        """Test parse_image_to_markdown successfully opens image and calls vision"""
        mock_img = MagicMock()
        mock_open.return_value.__enter__.return_value = mock_img
        mock_vision.return_value = "Parsed text from PNG screenshot."
        
        result = document_parser.parse_image_to_markdown("test_screenshot.png")
        self.assertIn("# Image Content: test_screenshot.png", result)
        self.assertIn("Parsed text from PNG screenshot.", result)
        mock_open.assert_called_once_with("test_screenshot.png")
        mock_vision.assert_called_once_with(mock_img)

if __name__ == '__main__':
    unittest.main()
