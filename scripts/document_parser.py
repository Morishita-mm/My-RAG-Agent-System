import os
import sys
import logging
import base64
import requests
from io import BytesIO
import dotenv
import re
import unicodedata

# 環境変数ロード
dotenv.load_dotenv()

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def encode_image_to_base64(image):
    buffered = BytesIO()
    # JPEG does not support alpha channel (transparency). Convert RGBA/P/LA to RGB.
    if image.mode in ('RGBA', 'P', 'LA'):
        image = image.convert('RGB')
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def analyze_image_with_vision(image, model_name=None) -> str:
    """Analyze page image using LiteLLM Vision API"""
    if model_name is None:
        model_name = os.environ.get("VISION_MODEL", "gemini-2.5-flash")
    base64_image = encode_image_to_base64(image)
    litellm_url = os.environ.get("LITELLM_API_BASE", "http://localhost:4000/v1")
    url = f"{litellm_url.rstrip('/')}/chat/completions"
    
    api_key = os.environ.get("LITELLM_API_KEY", "sk-1234")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    payload = {
        "model": model_name,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text", 
                        "text": "この画像のドキュメントページに含まれるテキスト、表、図表の内容をすべて読み取り、読みやすい構造化Markdown形式で出力してください。余計な説明や挨拶は一切含めず、出力されたMarkdownのみを返してください。"
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{base64_image}"
                        }
                    }
                ]
            }
        ]
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        if response.status_code == 200:
            res_data = response.json()
            return res_data.get("choices", [{}])[0].get("message", {}).get("content", "")
        else:
            logging.error(f"Vision API responded with {response.status_code}: {response.text}")
    except Exception as e:
        logging.error(f"Vision API connection failed: {e}")
    return ""

def parse_pdf_to_markdown(pdf_path: str, use_vision: bool = True) -> str:
    """Parse PDF to Markdown using PyPDF for text, and pdf2image + Vision for scanned/image pages"""
    from pypdf import PdfReader
    markdown_content = []
    
    try:
        reader = PdfReader(pdf_path)
        total_pages = len(reader.pages)
        logging.info(f"Parsing PDF '{pdf_path}' ({total_pages} pages)...")
        
        for page_idx, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            
            # テキスト文字数が極端に少ない場合は、スキャン画像または図表と判定してVisionを使用
            if use_vision and len(text.strip()) < 50:
                logging.info(f"Page {page_idx} has very little text ({len(text)} chars). Trying Vision OCR...")
                try:
                    from pdf2image import convert_from_path
                    # 特定ページのみを変換してメモリ負荷を軽減
                    images = convert_from_path(pdf_path, first_page=page_idx, last_page=page_idx)
                    if images:
                        vision_text = analyze_image_with_vision(images[0])
                        if vision_text:
                            markdown_content.append(f"## Page {page_idx} (Vision OCR Result)\n\n{vision_text}\n")
                            continue
                except Exception as e:
                    logging.warning(f"Failed to convert page {page_idx} to image (likely poppler is not installed): {e}")
            
            # テキストがある場合、または画像化に失敗した場合はそのままテキストを追加
            if text.strip():
                markdown_content.append(f"## Page {page_idx}\n\n{text}\n")
            else:
                markdown_content.append(f"## Page {page_idx}\n\n[Empty Page or Unparsed Content]\n")
                
    except Exception as e:
        logging.error(f"Error parsing PDF '{pdf_path}': {e}")
        return f"Error: Failed to parse PDF file: {e}"
        
    return "\n".join(markdown_content)

def parse_docx_to_markdown(docx_path: str) -> str:
    """Parse DOCX to Markdown (preserving paragraphs and tables)"""
    from docx import Document
    markdown_content = []
    
    try:
        doc = Document(docx_path)
        logging.info(f"Parsing DOCX '{docx_path}'...")
        
        # 1つのドキュメント内の段落とテーブルが混在するため、XML要素を順に走査する
        body_elements = doc.element.body
        paragraphs = {p.paragraph_format.element: p for p in doc.paragraphs}
        tables = {t._tbl: t for t in doc.tables}
        
        for child in body_elements:
            # 段落要素の処理
            if child in paragraphs:
                p = paragraphs[child]
                text = p.text.strip()
                if not text:
                    continue
                # 見出しスタイルの判定
                if p.style.name.startswith("Heading"):
                    try:
                        level = int(p.style.name.replace("Heading", "").strip())
                        markdown_content.append(f"{'#' * level} {text}\n")
                    except ValueError:
                        markdown_content.append(f"### {text}\n")
                else:
                    markdown_content.append(f"{text}\n")
                    
            # テーブル要素の処理
            elif child in tables:
                t = tables[child]
                markdown_content.append("\n")
                for row_idx, row in enumerate(t.rows):
                    row_cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    markdown_content.append(f"| {' | '.join(row_cells)} |")
                    # ヘッダー下のアライメント行を追加
                    if row_idx == 0:
                        align_row = ["---"] * len(row_cells)
                        markdown_content.append(f"| {' | '.join(align_row)} |")
                markdown_content.append("\n")
                
    except Exception as e:
        logging.error(f"Error parsing DOCX '{docx_path}': {e}")
        return f"Error: Failed to parse Word document: {e}"
        
    return "\n".join(markdown_content)

def parse_excel_to_markdown(excel_path: str) -> str:
    """Parse Excel to Markdown sheets"""
    import pandas as pd
    markdown_content = []
    
    try:
        logging.info(f"Parsing Excel '{excel_path}'...")
        xls = pd.ExcelFile(excel_path)
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            # 全ての値が空の列や行を排除
            df = df.dropna(how='all')
            if df.empty:
                continue
                
            markdown_content.append(f"## Sheet: {sheet_name}\n")
            # pandasのto_markdown()機能を使ってMarkdown表に高速変換
            markdown_table = df.to_markdown(index=False)
            markdown_content.append(markdown_table)
            markdown_content.append("\n")
            
    except Exception as e:
        logging.error(f"Error parsing Excel '{excel_path}': {e}")
        return f"Error: Failed to parse Excel document: {e}"
        
    return "\n".join(markdown_content)

def parse_image_to_markdown(image_path: str) -> str:
    """Parse direct image files (.png, .jpg, .jpeg) using Vision LLM"""
    from PIL import Image
    try:
        logging.info(f"Parsing Image '{image_path}' using Vision LLM...")
        with Image.open(image_path) as img:
            vision_text = analyze_image_with_vision(img)
            if vision_text:
                return f"# Image Content: {os.path.basename(image_path)}\n\n{vision_text}\n"
    except Exception as e:
        logging.error(f"Error parsing image '{image_path}': {e}")
        return f"Error: Failed to parse image: {e}"
    return "Error: Empty or unparseable image content."

def normalize_text(text: str) -> str:
    """Normalize text by converting full-width alphanumeric to half-width and cleaning up whitespaces"""
    if not text:
        return ""
    # NFKC正規化（全角英数字を半角に変換）
    normalized = unicodedata.normalize('NFKC', text)
    # 連続する複数のスペース（半角・全角・タブ）を1つの半角スペースに集約
    normalized = re.sub(r'[ \t\u3000]+', ' ', normalized)
    # 3つ以上連続する改行を2つに制限（段落の崩れ防止）
    normalized = re.sub(r'\n{3,}', '\n\n', normalized)
    return normalized

def generate_document_summary(markdown_text: str, model_name=None) -> str:
    """Generate a brief summary of the document using LiteLLM"""
    if not markdown_text.strip():
        return ""
    if model_name is None:
        model_name = os.environ.get("SUMMARY_MODEL", "gemini-2.5-flash")
    
    litellm_url = os.environ.get("LITELLM_API_BASE", "http://localhost:4000/v1")
    url = f"{litellm_url.rstrip('/')}/chat/completions"
    api_key = os.environ.get("LITELLM_API_KEY", "sk-1234")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    # 簡易カット（コンテキスト制限対策）
    truncated_text = markdown_text[:20000]
    
    prompt = """以下に示すドキュメントの内容を正確に理解し、このドキュメントが「何について書かれたものか」を日本語で3〜5文程度の簡潔な概要（サマリー）としてまとめてください。
余計な解説、前置き、挨拶は一切含めず、概要の文章のみを出力してください。

ドキュメント内容:
"""
    
    payload = {
        "model": model_name,
        "messages": [
            {
                "role": "user",
                "content": f"{prompt}\n\n{truncated_text}"
            }
        ],
        "temperature": 0.2
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=45)
        if response.status_code == 200:
            res_data = response.json()
            return res_data.get("choices", [{}])[0].get("message", {}).get("content", "").strip()
        else:
            logging.error(f"Summary API responded with {response.status_code}: {response.text}")
    except Exception as e:
        logging.error(f"Summary API connection failed: {e}")
    return ""

def convert_document_to_markdown(file_path: str, generate_summary: bool = False) -> str:
    """Main routing function to parse any supported document path to Markdown string"""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.pdf':
        markdown_str = parse_pdf_to_markdown(file_path)
    elif ext == '.docx':
        markdown_str = parse_docx_to_markdown(file_path)
    elif ext in ['.xlsx', '.xls', '.exls']:
        markdown_str = parse_excel_to_markdown(file_path)
    elif ext in ['.png', '.jpg', '.jpeg']:
        markdown_str = parse_image_to_markdown(file_path)
    elif ext == '.md':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                markdown_str = f.read()
        except Exception as e:
            markdown_str = f"Error reading markdown: {e}"
    else:
        markdown_str = f"Unsupported file extension: {ext}"

    if markdown_str.startswith("Error:"):
        return markdown_str

    # テキストの正規化
    markdown_str = normalize_text(markdown_str)

    # 必要に応じて要約の生成と追記
    if generate_summary:
        logging.info(f"Generating summary for document: {os.path.basename(file_path)}...")
        summary = generate_document_summary(markdown_str)
        if summary:
            summary_section = f"# Document Summary / ドキュメント概要\n\n{summary}\n\n---\n\n"
            markdown_str = summary_section + markdown_str

    return markdown_str
